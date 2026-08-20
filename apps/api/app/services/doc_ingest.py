"""Text extraction for knowledge document ingestion.

Turns uploaded files (PDF, Word, plain text/markdown) into plain text that
feeds the existing workspace-doc pipeline (`upsert_doc` -> chunk -> embed),
so uploaded documents become retrievable in agent drafts via hybrid search.
"""

from __future__ import annotations

import io
import re

from fastapi import HTTPException

# Extensions ingested by simple text decoding.
_TEXT_EXTENSIONS = {"txt", "md", "markdown", "csv", "tsv", "rst", "log", "json"}

SUPPORTED_EXTENSIONS = _TEXT_EXTENSIONS | {"pdf", "docx"}


def _extension(filename: str) -> str:
    return filename.rsplit(".", 1)[-1].lower() if "." in filename else ""


def _decode_text(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        return data.decode("latin-1", errors="replace")


def _extract_pdf(data: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise HTTPException(status_code=500, detail="PDF support is not installed") from exc
    try:
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise HTTPException(status_code=422, detail="Could not read this PDF file") from exc
    return "\n\n".join(p.strip() for p in pages if p.strip())


def _extract_docx(data: bytes) -> str:
    try:
        import docx
    except ImportError as exc:  # pragma: no cover - dependency is declared
        raise HTTPException(status_code=500, detail="Word support is not installed") from exc
    try:
        document = docx.Document(io.BytesIO(data))
    except Exception as exc:
        raise HTTPException(status_code=422, detail="Could not read this Word file") from exc
    parts: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n\n".join(parts)


def _normalize_whitespace(text: str) -> str:
    # Collapse runs of 3+ newlines; strip trailing spaces per line.
    text = "\n".join(line.rstrip() for line in text.splitlines())
    return re.sub(r"\n{3,}", "\n\n", text).strip()


def extract_text(filename: str, data: bytes) -> str:
    """Extract plain text from an uploaded document.

    Raises 415 for unsupported types and 422 when no text can be extracted
    (e.g. a scanned/image-only PDF).
    """
    ext = _extension(filename)
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(
            status_code=415,
            detail=(
                "Unsupported file type. Supported: "
                + ", ".join(sorted(SUPPORTED_EXTENSIONS))
            ),
        )
    if ext == "pdf":
        text = _extract_pdf(data)
    elif ext == "docx":
        text = _extract_docx(data)
    else:
        text = _decode_text(data)
    text = _normalize_whitespace(text)
    if not text:
        raise HTTPException(
            status_code=422,
            detail="No readable text found in this file (scanned documents are not supported yet).",
        )
    return text
